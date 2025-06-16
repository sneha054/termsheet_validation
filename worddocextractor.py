
from docx import Document

def extract_text_from_worddocx(docx_path):

    doc=Document(docx_path)

    full_text=[]

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_tables_from_worddocx(docx_path):
    doc=Document(docx_path)

    textundertables=[]

    for table in doc.tables:
        for row in table.rows:
            row_data=[cell.text.strip() for cell in row.cells]
            textundertables.append('\t'.join(row_data))
    return '\n'.join(textundertables)
    

docx_path='CFI-Term-Sheet-Template.docx'
text=extract_text_from_worddocx(docx_path)
text1=extract_tables_from_worddocx(docx_path)
print(text)
print(text1)
