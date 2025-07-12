from docx import Document
import sys

def extract_text_from_worddocx(docx_path):

    doc=Document(docx_path)

    full_text=[]

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_tables_from_worddocx(docx_path):
    doc=Document(docx_path)

    textundertables=[]

    for table in doc.tables:
        for row in table.rows:
            row_data=[cell.text.strip() for cell in row.cells]
            textundertables.append('\t'.join(row_data))
    return '\n'.join(textundertables)
    
def save_text_to_file(text:str,filename,mode):
    with open(filename,mode,encoding="utf-8") as f:
        f.write(text)

if __name__=="__main__":
    if len(sys.argv) < 2:
        print("Running word_extractor.py")
        sys.exit(1)

    docx_path=sys.argv[1]
    text=extract_text_from_worddocx(docx_path)
    text1=extract_tables_from_worddocx(docx_path)
    print(text)
    print(text1)
    output_path="D:\\termsheet_validation\\docx_extraction\\text_output2.txt"
    save_text_to_file(text,output_path,'w')
    save_text_to_file(text1,output_path,'a')
